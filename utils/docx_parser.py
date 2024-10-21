from docx import Document
import re

PREGUNTA_KEYWORD = "pregunta"
ROTADA_KEYWORD = "rotada"
PRESIONE_KEYWORD = "presione"
SERIE_TEMPORAL_KEYWORD = 'serie temporal'
def es_pregunta_valida(texto):
    return (
        PREGUNTA_KEYWORD in texto.lower() 
        and 'ahora' not in texto.lower()
        and 'y si' not in texto.lower() 
        and 'quisiéramos' not in texto.lower()
    )

def extraer_numero_pregunta(texto):
    match = re.search(r'(\d+)', texto)
    if match:
        return 'P' + str.zfill(match.group(1), 2)
    return None

def get_preguntas(ruta):
    archivo = Document(ruta)
    preguntas = {}
    codigo = None
    banderin_apertura = False

    for parrafo in archivo.paragraphs:
        texto = parrafo.text.strip()

        if not texto:
            continue
        
        if ROTADA_KEYWORD in texto.lower():
            continue
        
        if es_pregunta_valida(texto):
            codigo = extraer_numero_pregunta(texto)
            if codigo:
                preguntas[codigo] = ''
                banderin_apertura = True
            continue
        
        if codigo and banderin_apertura and PRESIONE_KEYWORD not in texto.lower() and SERIE_TEMPORAL_KEYWORD not in texto.lower():
            preguntas[codigo] = texto
            banderin_apertura = False
            
            if not preguntas[codigo]:
                print(f"Advertencia: La pregunta con código '{codigo}' tiene un valor vacío.")
    
    return preguntas

def getOpciones(ruta):
    archivo = Document(ruta)
    opciones = {}
    for parrafo in archivo.paragraphs:
        texto = parrafo.text
        if not texto:
            continue
        if ROTADA_KEYWORD in texto.lower():
            continue
        if 'Pregunta' in texto:
            codigo = 'P' + str.zfill(texto.strip('Pregunta '), 2)
            opciones[codigo] = []
        
        if PRESIONE_KEYWORD in texto:
            opciones[codigo].append(texto)
    
    opciones_limpias = {pregunta: [opcion.replace(" Presione ", "").replace(" presione ", "").replace('.', '')[:-2] for opcion in opciones[pregunta]]
                        for pregunta in opciones}
    
    return opciones_limpias

def getBloque(ruta):
    archivo = Document(ruta)
    bloques_dict = {}
    bloque = ''
    for parrafo in archivo.paragraphs:
        texto = parrafo.text
        
        if not texto:
            continue
        
        if 'bloque' in texto.lower():
            bloque = texto.replace('Bloque ', '').capitalize()
        elif 'reach' in texto.lower():
            bloque = 'Reach de intención de voto'
            
        if 'pregunta' in texto.lower():
            codigo = 'P' + str.zfill(texto.strip('Pregunta '), 2)
            bloques_dict[codigo] = bloque
            
    return bloques_dict