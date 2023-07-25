# -*- coding: utf-8 -*-
"""
Spyder Editor

This is a temporary script file.
"""

from docx import Document
import pandas as pd
import seaborn as sns
import os
import csv

def extractorPreguntas(preguntas, texto, banderin_apertura, codigo):    

     if not texto:
         return
     if 'rotada' in texto.lower() or 'rotadas' in texto.lower():
         return
     # bis =  'BIS' in texto.lower() or 'bis' in texto.lower()
     
     if 'pregunta' in texto.lower() and 'ahora' not in texto.lower():
         preguntas[codigo] = ''
         banderin_apertura = True
         return banderin_apertura
     
     elif codigo in preguntas.keys() and banderin_apertura == True:
         
         if 'Presione' not in texto:
             preguntas[codigo]= texto
             banderin_apertura = False    
             return(preguntas, banderin_apertura)
     return(preguntas)

ruta = ruta_cuestionario
def getTodo(ruta):
    ''' 
    Genera un diccionario con el siguiente formato
    {'P01':['opcion_1', ... ,'opcion_N'],
     'P02':['opcion_1', ... ,'opcion_N'],
     etc}
    y lo printea en consola
    
    Excepciones: Preguntas rotadas y preguntas BIS
    Las opciones las encuentra a partir de el texto "presione" 
    '''
    archivo = Document(ruta)
    opciones = {}
    preguntas = {}
    bloques_dict = {}
    
    bloque = ''
    codigo = 0

    banderin_apertura = False
    
    for parrafo in archivo.paragraphs:

        texto = parrafo.text
        
        if not texto:
            continue
        
        #### Bloques
        
        if 'bloque' in texto.lower():
            bloque = texto.replace('Bloque ', '').capitalize()
     
        
        #### Preguntas
        if 'rotada' in texto.lower() or 'rotadas' in texto.lower():
            continue
        elif 'Pregunta' in texto:
            codigo = texto.strip('Pregunta ')
            codigo = 'P'+str.zfill(codigo, 2)
            opciones[codigo] = []
            bloques_dict[codigo] = bloque
            if 'ahora' not in texto.lower():
                preguntas[codigo] = ''
                banderin_apertura = True
                continue
            
        elif 'Presione' in texto or 'presione' in texto:
            opciones[codigo].append(texto)
        
        
        elif codigo in preguntas.keys() and banderin_apertura == True:
            
            if 'Presione' not in texto:
                preguntas[codigo]= texto
                banderin_apertura = False     


    opciones_limpias = {}
    for pregunta in opciones:
        opciones_limpias[pregunta] = []
        if pregunta == "P1":
            opciones_limpias["P1"] = opciones["P1"]
        
        for indice, opcion in enumerate(opciones[pregunta]):
            texto_limpio = opcion.replace(" Presione ", "")[:-2]
            texto_limpio = texto_limpio.replace('.','')
            opciones_limpias[pregunta].append(texto_limpio)
    
    for key in opciones_limpias:
        if 'BIS' not in str(key):
            print(f"'{key}' : ", opciones_limpias[key])
            print()
    
    for key in preguntas:
        if 'BIS' not in str(key) and 'bis' not in str(key):
            print(f"'{key}' : '{preguntas[key].strip()}'")
            print()
    
            


    for parrafo in archivo.paragraphs:
        texto = parrafo.text
        
        if not texto:
            continue
        
    return(opciones_limpias, preguntas, bloques_dict)

def infer_politica(opcion):

    dic = dict(
    peronismo = ['juntos avancemos','por la patria','nos une','peron','peronismo', 'frente de todos', 'fdt','evita', 'kirchnerismo', 'alberto fernandez', 'brutten','kirchner', 'cfk','vamos con todos'],
    cambiemos = ['unidos para cambiar', 'macrismo', 'cambiemos', 'juntos por el cambio', 'junto x el cambio', 'pro', 'jxc','tortoriello'],
    liberales = ['lieberal', 'libertar', 'milei', 'libertad'],
    izquierda = ['izquierda', 'socialismo'],
    blanco = ['en blanco'],
    nosabe = ['no sabe', 'no se'])
    preonismo_nok = ['randazzo', 'hacemos por nuestro']

    for color, palabras in dic.items():
        for palabra in palabras:
            if palabra in opcion.lower():
                return color

    return 'otros'

def list2dictPalette(respuestas, palette="RdYlGn_r", noSabe=True):
    if noSabe:
        # 'No sabe' tiene que estar en último lugar en la lista de respuestas
        colors = sns.color_palette(palette, len(respuestas) - 1)
        colors = colors.as_hex()
        # Cambio el color del amarillo por uno más fuerte
        colors.append('#c9bebd')
    else:
        # para preguntas que no tienen opción 'No sabe'
        colors = sns.color_palette(palette, len(respuestas))
        colors = colors.as_hex()
        # Cambio el color del amarillo por uno más fuerte
    colors = ['#fffd5e' if c == '#fffebe' else c for c in colors]
    return dict(zip(respuestas, colors))

def infer_paleta(opciones_todas, cuestionario, default_paleta="qualitative_strong"):

    # import pandas as pd

    # path = path[:-1] if path.endswith('/') else path

    df_preguntas = cuestionario

    paleta = {}
    bloque = ''
    for i, row in df_preguntas.fillna("").iterrows():
        pregunta = row.codigo
        paleta = row.paleta
        codigo = row.codigo
        label = row.label
        if bloque != row.bloque:
            bloque = row.bloque
            print()
            print(f'# Bloque {bloque}')
        if row.paleta == "imagen":
            print(f'paletas["{pregunta}"] = opinionColorDict')


        else:

            opciones = opciones_todas[pregunta]

            if "Buena" in opciones:
                # if row.paleta == "imagen":
                print(f'paletas["{pregunta}"] = opinionColorDict')
                
            elif 'votaría' in opciones or paleta in ["votaria", 'reach']:
                print(f'paletas["{pregunta}"] = votaria')


            elif paleta == "politica":
                apariciones = {'peronismo':0, 'cambiemos':0, 'liberales':0, 'izquierda':0}
                p = {}
                for opcion in opciones:
                    politica = infer_politica(opcion)
                    # print(politica)
                    # print(opcion)
                    if politica in list(apariciones):
                        apariciones[politica]+= 1

                        if apariciones[politica]>1:
                            p[opcion] = f'colores["{infer_politica(opcion)}{apariciones[politica]}"]'
                        else:
                            p[opcion] = f'colores["{infer_politica(opcion)}"]'
                    else:
                        p[opcion] = f'colores["{infer_politica(opcion)}"]'
                    # print(p)
                # fix comillas molestas 
                value = repr(p).replace("'colores","colores").replace("]'", "]")

                print(f'paletas["{pregunta}"] = ' + value + '\n', sep = '\n')
                
            elif paleta.lower() == 'sino':
                print(f'paletas["{pregunta}"] = siNoColorDict')
                

            elif paleta.lower() == 'frecuencia':
                print(f'paletas["{pregunta}"] = frecuenciaColorDict')
            
            elif 'mucho' in paleta.lower():
                print(f'paletas["{pregunta}"] = muchoPocoColorDict')
                
            elif 'sino' in paleta.lower():
                print(f'paletas["{pregunta}"] = siNoColordict')
            else:

                opciones_lower = [ ]
                for opcion in opciones:
                    opciones_lower.append(opcion.lower())
                    
                nosabe =  ("no sabe" in opciones_lower )  or ( "no se" in opciones_lower)

                print(f'paletas["{pregunta}"] = list2dictPalette({repr(opciones)}, {default_paleta}, noSabe={nosabe})')

def cuestionarioTSV(rutaCuestionario, rutaGuardado='/home/julian/trabajo/updates', nombreArchivo = ''):
    '''
    Lee las preguntas y bloques previamente pasados a diccionario y arma un csv
    completando código, bloque y texto
    '''
    preguntas = getPreguntas(rutaCuestionario)
    bloques = getBloque(rutaCuestionario)
    claves = preguntas.keys()
    campos =['codigo','bloque','paleta','texto','label','aclaracion']
    
    
    while len(nombreArchivo)<1:
        nombreArchivo = input('nombre del tsv: ')
    
    if '.tsv' not in nombreArchivo:
        nombreArchivo += '.tsv'
    
    rutaCompleta = rutaGuardado+'/'+nombreArchivo        
    with open(rutaCompleta, 'w', newline='') as archivoTSV:
        writer = csv.writer(archivoTSV, delimiter='\t')
        
        writer.writerow(campos)

        # Iterar sobre las claves y escribir los valores correspondientes en cada columna
        for clave in claves:
            
            pregunta = preguntas[clave]
            bloque = bloques[clave]
            if 'BIS' in clave or 'bis' in clave:
                continue
            # Escribir una nueva fila en el archivo TSV
            writer.writerow([clave, bloque, '', pregunta, '', ''])
    
        print('Archivo TSV generado exitosamente.')


