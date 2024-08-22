# -*- coding: utf-8 -*-
"""
Spyder Editor

This is a temporary script file.
"""

from docx import Document
import pandas as pd
import seaborn as sns
import os
import csv
import re
def getOpciones(ruta):
    ''' 
    Genera un diccionario con el siguiente formato
    {'P01':['opcion_1', ... ,'opcion_N'],
     'P02':['opcion_1', ... ,'opcion_N'],
     etc}
    y lo printea en consola
    
    Excepciones: Preguntas rotadas y preguntas BIS
    Las opciones las encuentra a partir de el texto "presione" 
    '''
    archivo = Document(ruta)
    opciones = {}
    for parrafo in archivo.paragraphs:
        texto = parrafo.text
        if not texto:
            continue
        if 'rotada' in texto.lower() or 'rotadas' in texto.lower() or 'rotacion' in texto.lower():
            continue
        if 'Pregunta' in texto:
            codigo = texto.strip('Pregunta ')
            codigo = 'P'+str.zfill(codigo, 2)
            opciones[codigo] = []
        
        if 'Presione' in texto or 'presione' in texto:
            opciones[codigo].append(texto)
    
    opciones_limpias = {}
    for pregunta in opciones:
        opciones_limpias[pregunta] = []
        if pregunta == "P1":
            opciones_limpias["P1"] = opciones["P1"]
        
        for indice, opcion in enumerate(opciones[pregunta]):
            texto_limpio = opcion.replace(" Presione ", "").replace(" presione ", "")[:-2]
            texto_limpio = texto_limpio.replace('.','')
            opciones_limpias[pregunta].append(texto_limpio)
    
    for key in opciones_limpias:
        if 'BIS' not in str(key):
            print(f"'{key}' : ", opciones_limpias[key])
            print()
            
    return(opciones_limpias)

PREGUNTA_KEYWORD = "pregunta"
ROTADA_KEYWORD = "rotada"
PRESIONE_KEYWORD = "presione"

def es_pregunta_valida(texto):
    """Verifica si el texto corresponde a una pregunta válida."""
    return (
        PREGUNTA_KEYWORD in texto.lower() 
        and 'ahora' not in texto.lower()
        and 'y si' not in texto.lower() 
        and 'quisiéramos' not in texto.lower()
    )

def extraer_numero_pregunta(texto):
    """Extrae el número de la pregunta del texto."""
    match = re.search(r'(\d+)', texto)
    if match:
        return 'P' + str.zfill(match.group(1), 2)
    return None

def get_preguntas(ruta):
    archivo = Document(ruta)
    preguntas = {}
    codigo = None
    banderin_apertura = False

    for parrafo in archivo.paragraphs:
        texto = parrafo.text.strip()

        if not texto:
            continue
        
        if ROTADA_KEYWORD in texto.lower():
            continue
        
        if es_pregunta_valida(texto):
            codigo = extraer_numero_pregunta(texto)
            if codigo:
                preguntas[codigo] = ''
                banderin_apertura = True
            continue
        
        if codigo and banderin_apertura and PRESIONE_KEYWORD not in texto.lower():
            preguntas[codigo] = texto
            banderin_apertura = False
            
            # Verificar si el valor es vacío y emitir una advertencia
            if not preguntas[codigo]:
                print(f"Advertencia: La pregunta con código '{codigo}' tiene un valor vacío.")
    
    # Impresión del resultado
    for key, value in preguntas.items():
        if 'BIS' not in key:
            print(f"'{key}' : '{value.strip()}'\n")
        
    return preguntas

def getBloque(ruta):
    ''' 
    Genera un diccionario con el siguiente formato
    {'P01':'bloque'
     'P02':,
     etc}
    y lo printea en consola
    '''
    archivo = Document(ruta)
    bloques_dict = {}
    bloque = ''
    for parrafo in archivo.paragraphs:
        texto = parrafo.text
        
        if not texto:
            continue
        
        if 'bloque' in texto.lower():
            bloque = texto.replace('Bloque ', '').capitalize()
        elif 'reach' in texto.lower():
            bloque = 'Reach de intención de voto'
        
        # if bloque.lower() == 'sociodemográfico':
        #     bloque = 'Control'
            
        if 'pregunta' in texto.lower():
            codigo = texto.strip('Pregunta ')
            codigo = 'P'+str.zfill(codigo, 2)
            bloques_dict[codigo] = bloque
    return(bloques_dict)


def infer_politica(opcion):

    dic = dict(
    peronismo = ['massa', 'kicillof','juntos avancemos','por la patria','nos une','peron','peronismo', 'frente de todos', 'fdt','evita', 'kirchnerismo', 'alberto fernandez', 'brutten','kirchner', 'cfk','vamos con todos'],
    cambiemos = ['encuentro por corrientes', 'bullrich','unidos para cambiar', 'macrismo', 'cambiemos', 'juntos por el cambio', 'junto x el cambio', 'pro', 'jxc','tortoriello'],
    liberales = ['lieberal', 'libertar', 'milei', 'libertad', 'villarruel'],
    izquierda = ['izquierda', 'socialismo', 'bregman'],
    peronismo_nok = ['país', 'pais', 'randazzo', 'schiaretti'],
    blanco = ['en blanco'],
    nosabe = ['no sabe', 'no se'])

    for color, palabras in dic.items():
        for palabra in palabras:
            if palabra in opcion.lower():
                return color

    return 'otros'

def list2dictPalette(respuestas, palette="RdYlGn_r", noSabe=True):
    if noSabe:
        # 'No sabe' tiene que estar en último lugar en la lista de respuestas
        colors = sns.color_palette(palette, len(respuestas) - 1)
        colors = colors.as_hex()
        # Cambio el color del amarillo por uno más fuerte
        colors.append('#c9bebd')
    else:
        # para preguntas que no tienen opción 'No sabe'
        colors = sns.color_palette(palette, len(respuestas))
        colors = colors.as_hex()
        # Cambio el color del amarillo por uno más fuerte
    colors = ['#fffd5e' if c == '#fffebe' else c for c in colors]
    return dict(zip(respuestas, colors))

def infer_paleta_else(paleta):
    if 'sino' == paleta:
        return('paletas["{pregunta}"] = siNoColorDict')
    
    elif 'sinonosabe' in paleta:
        return('paletas["{pregunta}"] = siNoNoSabeColorDict')
        
    elif paleta in ['frecuencia', 'freq', 'frec']:
        return('paletas["{pregunta}"] = frecuenciaColorDict')
    
    elif 'mucho' in paleta:
        return('paletas["{pregunta}"] = muchoPocoColorDict')
    
    elif paleta in ['problema', 'deuda', 'ingresos', 'opinion', 'comparacion', 'medios']:
        return('paletas["{pregunta}"] = {paleta}ColorDict')
    else:
        return False
    

def infer_paleta(opciones_todas, cuestionario):
    
    df_preguntas = cuestionario

    paleta = {}
    bloque = ''
    for i, row in df_preguntas.fillna("").iterrows():
        default_paleta="diverging"
        pregunta = row.codigo
        paleta = row.paleta
        codigo = row.codigo
        # if not pregunta[1].isdigit() or 'imputada' in pregunta:
        #     continue
        if bloque != row.bloque:
            bloque = row.bloque
            print()
            print(f'# Bloque {bloque}')
        if row.paleta == "imagen":
            print(f'paletas["{pregunta}"] = opinionColorDict')


        else:

            opciones = opciones_todas[pregunta]

            if "Buena" in opciones:
                # if row.paleta == "imagen":
                print(f'paletas["{pregunta}"] = opinionColorDict')
                
            elif 'votaría' in opciones or paleta in ["votaria", 'reach']:
                print(f'paletas["{pregunta}"] = votaria')


            elif paleta == "politica":
                apariciones = {'peronismo':0, 'cambiemos':0, 'liberales':0, 'izquierda':0}
                p = {}
                for opcion in opciones:
                    politica = infer_politica(opcion)
                    if politica in list(apariciones):
                        apariciones[politica]+= 1

                        if apariciones[politica]>1:
                            p[opcion] = f'colores["{infer_politica(opcion)}{apariciones[politica]}"]'
                        else:
                            p[opcion] = f'colores["{infer_politica(opcion)}"]'
                    else:
                        p[opcion] = f'colores["{infer_politica(opcion)}"]'

                # fix comillas molestas 
                value = repr(p).replace("'colores","colores").replace("]'", "]")

                print(f'paletas["{pregunta}"] = ' + value + '\n', sep = '\n')
            
            else:
                if infer_paleta_else(paleta.lower()):
                    string_salida = infer_paleta_else(paleta.lower()).format(pregunta = pregunta, paleta = paleta.lower())
                    print(string_salida)
                else:
                    opciones_lower = [ ]
                    for opcion in opciones:
                        opciones_lower.append(opcion.lower())
                        
                    nosabe =  ("no sabe" in opciones_lower )  or ( "no se" in opciones_lower)
                    if 'quali' in paleta:
                        default_paleta = 'qualitative_strong'
                        
                    print(f'paletas["{pregunta}"] = list2dictPalette({repr(opciones)}, {default_paleta}, noSabe={nosabe})')

def limpiarKeys(diccionario):
    
    diccionarioLimpio = {}
    for key in diccionario:
        if len(key) != 3 and key not in diccionarioLimpio:
            keyLimpia = key[:3]
            if keyLimpia[-1] == ' ':
                numero = keyLimpia.strip('P ')
                keyLimpia = 'P'+str.zfill(numero, 2)
        else:
            keyLimpia = key
            
        if diccionario[key]:
            diccionarioLimpio[keyLimpia] = diccionario[key]
    
    return diccionarioLimpio

def cuestionarioTSV(rutaCuestionario, rutaGuardado='/home/julian/trabajo/updates', nombreArchivo = '', 
                    checkCarpeta = True, rutaData = 'web/individuales'):
    '''
    Lee las preguntas y bloques previamente pasados a diccionario y arma un csv
    completando código, bloque y texto
    '''
    rutaCompleta = os.path.join(rutaGuardado,nombreArchivo)
    preguntas = limpiarKeys(get_preguntas(rutaCuestionario))
    claves = preguntas.keys()
    
    bloques = limpiarKeys(getBloque(rutaCuestionario)) or dict(zip(claves, ['General']*len(claves))) #por si no hay ningún bloque marcadoe n la encuesta
    
    campos =['codigo','bloque','paleta','texto','label','aclaracion']
    
    rutaIndividuales = os.path.join(rutaGuardado,rutaData)
    if checkCarpeta and os.path.exists(rutaIndividuales):
        claves = sorted([key.lstrip('pregunta_').rstrip('.csv') for key in os.listdir(rutaIndividuales)])
    
    while len(nombreArchivo)<1:
        nombreArchivo = input('nombre del tsv: ')
    
    if '.tsv' not in nombreArchivo:
        nombreArchivo += '.tsv'
    
        
    with open(rutaCompleta, 'w', newline='') as archivoTSV:
        writer = csv.writer(archivoTSV, delimiter='\t')
        writer.writerow(campos)

        # Iterar sobre las claves y escribir los valores correspondientes en cada columna
        for clave in claves:
            
            if 'BIS' in clave or 'bis' in clave:
                continue
            if clave.startswith('P') and clave[1].isdigit():
                alcaracion = ''
                if 'imputada' not in clave and 'repregunta' not in clave:
                    pregunta = preguntas[clave]
                    bloque = bloques[clave]
                
                elif 'imputada' in clave:
                    pregunta = preguntas[clave.replace('_imputada', '')]
                    bloque = bloques[clave.replace('_imputada', '')]
                    aclaracion = '(Imputación no lineal de indecisos)'
                elif 'repregunta' in clave:
                    pregunta = preguntas[clave.replace('_repregunta', '')]
                    bloque = bloques[clave.replace('_repregunta', '')]
                    aclaracion = '(Repregunta)'
            else:
                bloque = ''
                pregunta = ''
            linea = [clave, bloque, '', pregunta, '', alcaracion]
            writer.writerow(linea)      
        print('Archivo TSV generado exitosamente.')

def getOpcionesFromData(ruta, subcarpeta = 'web/individuales'):
    individuales = os.listdir(os.path.join(ruta, subcarpeta))
    opciones = {}
    for archivo in individuales:
        df = pd.read_csv(os.path.join(ruta, subcarpeta, archivo), delimiter = ',')
        opciones_individual = list(df.clase)
        key = archivo.lstrip('pregunta_').rstrip('.csv')
        opciones[key] = opciones_individual
    return opciones
        
def printPaletas(texto, nombre_tsv, ruta_trabajo, subcarpeta = 'web/individuales'):
    ruta_cuestionario = ruta_trabajo + texto
    cuestionario = pd.read_table(ruta_trabajo + nombre_tsv)
    opciones = getOpcionesFromData(ruta_trabajo, subcarpeta)
    
    opciones_limpias = {}
    for key in opciones:
        codigo_limpio = key.replace('.csv', '')
        if codigo_limpio in cuestionario.codigo.to_list():
            opciones_limpias[codigo_limpio] = opciones[key]
    
    
    
    bloques = cuestionario.groupby('bloque', sort=False).agg(list)['codigo'].to_dict()
    
    keys_archivos = [k.lstrip('pregunta_').rstrip('.csv') for k in os.listdir(os.path.join(ruta_trabajo, subcarpeta))]
    for key in cuestionario.codigo.to_list():
        if key not in keys_archivos:
            cuestionario = cuestionario[cuestionario['codigo']!= key]
            print(f'{key} no está en el directorio')
    infer_paleta(opciones_limpias, cuestionario)


    return(opciones_limpias)